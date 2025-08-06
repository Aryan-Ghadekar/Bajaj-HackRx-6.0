import asyncio, hashlib, io, json, os, pickle, re, datetime, pathlib
from typing import List, Optional
from fastapi.staticfiles import StaticFiles
import aiohttp, numpy as np, tiktoken, faiss
from fastapi import FastAPI, UploadFile, File, Form, Header, HTTPException
from pydantic import BaseModel
from sentence_transformers import SentenceTransformer, CrossEncoder
from dotenv import load_dotenv
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from pathlib import Path

from database import Session, engine
from models   import Data

# ──────────────────────────── static folder path ──────────────────────────────
STATIC_FOLDER = Path(r"C:\Users\hp\OneDrive\Desktop\Bajaj\static")

# ──────────────────────────── NLTK setup ─────────────────────────────────────
import nltk, warnings, re as _re
from contextlib import suppress

for res in ("punkt_tab", "punkt"):
    try:
        nltk.data.find(f"tokenizers/{res}")
        break
    except LookupError:
        with suppress(Exception):
            nltk.download(res, quiet=True)
            try:
                nltk.data.find(f"tokenizers/{res}")
                break
            except LookupError:
                pass
else:
    warnings.warn("NLTK models missing; using simple regex sentence splitter")

    def _simple_tok(txt: str):
        return [s.strip() for s in _re.split(r"(?<=[.!?])\s+", txt) if s.strip()]

    nltk.sent_tokenize = _simple_tok
# ─────────────────────────────────────────────────────────────────────────────

load_dotenv()

EMBED_MODEL  = "BAAI/bge-base-en-v1.5"
CROSS_ENC    = "cross-encoder/ms-marco-MiniLM-L-6-v2"

AUTH_TOKEN   = "13e8f1be06f08115af28397a9f74cd278d9fc81b65945ec8a1c19a2c45511a4c"
GEMINI_KEY   = os.getenv("GEMINI_API_KEY")
GROQ_KEY     = os.getenv("GROQ_API_KEY")
GEMINI_URL   = (
    "https://generativelanguage.googleapis.com/v1beta/"
    "models/gemini-1.5-flash-latest:generateContent"
)

for var, val in (("GEMINI_API_KEY", GEMINI_KEY), ("GROQ_API_KEY", GROQ_KEY)):
    if not val:
        raise ValueError(f"{var} environment variable is not set")

print("Loading embedding and cross-encoder models …")
embedder      = SentenceTransformer(EMBED_MODEL)
cross_encoder = CrossEncoder(CROSS_ENC)

# ──────────────────────────── DB session ─────────────────────────────────────
session=Session(bind=engine)

# ───────────────────────────── FastAPI ───────────────────────────────────────
app = FastAPI()


# class QueryRequest(BaseModel):
#     documents: str
#     questions: List[str]


class QueryResponse(BaseModel):
    answers: List[str]

def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(pdf_bytes))
    return "\n".join([page.extract_text() or "" for page in reader.pages])

def extract_text_from_docx(docx_bytes: bytes) -> str:
    doc = DocxDocument(io.BytesIO(docx_bytes))
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip() != ""])


ENCODER   = tiktoken.encoding_for_model("text-embedding-3-small")
CACHE_DIR = pathlib.Path("./embed_cache")
CACHE_DIR.mkdir(exist_ok=True)


def _md5(txt: str) -> str:
    return hashlib.md5(txt.encode()).hexdigest()


async def _download(url: str) -> bytes:
    async with aiohttp.ClientSession() as s:
        async with s.get(url, timeout=60) as r:
            if r.status != 200:
                raise HTTPException(400, f"Fetch error {r.status}")
            return await r.read()


def _clean(text: str) -> str:
    seen, out = set(), []
    for line in text.splitlines():
        t = line.strip()
        if not t or re.match(r"^(page\s*)?\d+$", t.lower()) or t.lower() in seen:
            continue
        seen.add(t.lower())
        out.append(line)
    txt = "\n".join(out)
    txt = re.sub(r"\n{2,}", "\n\n", txt)
    return re.sub(r"[ \t]{2,}", " ", txt)


def _extract(blob: bytes, ext: str) -> str:
    if ext == ".pdf":
        reader = PdfReader(io.BytesIO(blob))
        raw    = "\n".join(p.extract_text() or "" for p in reader.pages)
    else:
        doc = DocxDocument(io.BytesIO(blob))
        raw = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return _clean(raw)


async def load_document(url: str) -> str:
    ext = os.path.splitext(url)[1].lower()
    if ext not in {".pdf", ".docx"}:
        raise HTTPException(400, "Unsupported file type (PDF/DOCX).")
    data = await _download(url)
    return _extract(data, ext)


def sentence_chunk(text: str, max_tokens: int = 512, overlap_pct: float = 0.1) -> List[str]:
    sents           = nltk.sent_tokenize(text)
    chunks          = []
    buf, buf_tokens = [], 0
    overlap_tokens  = int(max_tokens * overlap_pct)

    for sent in sents:
        tok_cnt = len(ENCODER.encode(sent))
        if buf_tokens + tok_cnt > max_tokens and buf:
            chunks.append(" ".join(buf))
            # keep last sentences worth `overlap_tokens`
            while buf and buf_tokens > overlap_tokens:
                buf_tokens -= len(ENCODER.encode(buf[0]))
                buf.pop(0)
        buf.append(sent)
        buf_tokens += tok_cnt

    if buf:
        chunks.append(" ".join(buf))
    return chunks


def cached_embeddings(doc_hash: str, chunks: List[str]) -> np.ndarray:
    fp = CACHE_DIR / f"{doc_hash}.pkl"
    if fp.exists():
        return pickle.loads(fp.read_bytes())
    embeds = embedder.encode(chunks, batch_size=32, normalize_embeddings=True)
    fp.write_bytes(pickle.dumps(np.array(embeds)))
    return embeds


def _faiss_index(dim: int) -> faiss.Index:
    idx = faiss.IndexFlatIP(dim)
    if faiss.get_num_gpus() > 0:
        res = faiss.StandardGpuResources()
        idx = faiss.index_cpu_to_gpu(res, 0, idx)
    return idx


def faiss_search(query: str, texts: List[str], embeds: np.ndarray, top_k: int = 8):
    q_emb  = embedder.encode([query], normalize_embeddings=True)
    index  = _faiss_index(embeds.shape[1])
    index.add(embeds.astype(np.float32))
    scores, idxs = index.search(q_emb.astype(np.float32), top_k)
    return [{"text": texts[i], "score": float(scores[0][r])} for r, i in enumerate(idxs[0])]


def local_rerank(query: str, passages: List[dict]) -> List[dict]:
    pairs  = [[query, p["text"]] for p in passages]
    scores = cross_encoder.predict(pairs)
    for s, p in zip(scores, passages):
        p["rerank"] = float(s)
    return sorted(passages, key=lambda x: x["rerank"], reverse=True)


async def gemini_rerank(http, query: str, passages: List[dict]) -> List[dict]:
    joined = "\n\n".join(f"Passage {i+1}: {p['text']}" for i, p in enumerate(passages))
    prompt = (
        f"Query: {query}\n\nRank the passages by relevance. "
        f"Return a JSON list of passage numbers.\n\n{joined}"
    )
    payload = {
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {"temperature": 0, "maxOutputTokens": 60},
    }
    async with http.post(
        GEMINI_URL,
        headers={"Content-Type": "application/json"},
        params={"key": GEMINI_KEY},
        json=payload,
        timeout=15,
    ) as r:
        if r.status != 200:
            return passages
        try:
            order = json.loads((await r.json())["candidates"][0]["content"]["parts"][0]["text"])
            return [passages[i - 1] for i in order if 1 <= i <= len(passages)]
        except Exception:
            return passages


async def groq_answer(http, question: str, ctx: List[str]) -> str:
    prompt = (
        "You are an expert legal/insurance document analyst.\n"
        "Answer the user's question using ONLY the context below. "
        "Cite clause numbers in your rationale.\n\n"
        f"Context:\n{'\n\n'.join(ctx)}\n\nQuestion: {question}"
    )
    payload = {
        "model": "llama3-8b-8192",
        "messages": [
            {"role": "system", "content": "You are an expert legal/insurance document analyst."},
            {"role": "user",   "content": prompt},
        ],
        "temperature": 0,
        "max_tokens": 350,
    }
    async with http.post(
        "https://api.groq.com/openai/v1/chat/completions",
        headers={"Authorization": f"Bearer {GROQ_KEY}", "Content-Type": "application/json"},
        json=payload,
        timeout=30,
    ) as r:
        if r.status != 200:
            raise HTTPException(502, f"Groq error {r.status}")
        return (await r.json())["choices"][0]["message"]["content"].strip()


def trim_context(txt: str, max_words: int = 100) -> str:
    words = txt.split()
    return " ".join(words[:max_words]) if len(words) > max_words else txt

# ───────────────────────────── Sample Pdf ────────────────────────────────────
app.mount("/static", StaticFiles(directory="static"), name="static")

# ───────────────────────────── API ROUTE ────────────────────────────────────
@app.post("/api/v1/hackrx/run", response_model=QueryResponse)
async def run_query(document: UploadFile = File(...),
    questions: List[str] = Form(...), 
    authorization: Optional[str] = Header(None)):
    if authorization is None or not authorization.startswith("Bearer "):
        raise HTTPException(401, "Invalid or missing auth token.")

    token = authorization.split("Bearer ")[1]
    if token != AUTH_TOKEN:
        raise HTTPException(401, "Invalid or missing auth token.")

    # === Save file to static folder ===
    file_location =STATIC_FOLDER / document.filename
    with open(file_location, "wb") as f:
        f.write(await document.read())

    # === Proceed with reading for processing ===
    if document.filename.lower().endswith('.pdf'):
        with open(file_location, "rb") as f:
            doc_text = extract_text_from_pdf(f.read())
    elif document.filename.lower().endswith('.docx'):
        with open(file_location, "rb") as f:
            doc_text = extract_text_from_docx(f.read())
    else:
        raise HTTPException(status_code=400, detail="Unsupported file format")

    # 1. Prepare document
    doc_txt  = doc_text
    doc_hash = _md5(doc_txt)
    chunks   = sentence_chunk(doc_txt)
    embeds   = cached_embeddings(doc_hash, chunks)

    # 2. Answer each question
    async with aiohttp.ClientSession() as http:
        async def handle(q: str) -> str:
            passages  = faiss_search(q, chunks, embeds, top_k=8)
            passages  = local_rerank(q, passages)
            if passages[0]["rerank"] - passages[-1]["rerank"] < 0.05:
                passages = await gemini_rerank(http, q, passages)

            ctx = [trim_context(p["text"], 120) for p in passages[:4]]
            ans = await groq_answer(http, q, ctx)

            if len(ans) < 50 or any(k in ans.lower() for k in ("i don't know", "no relevant information", "unable to answer")):
                ctx_long = [trim_context(p["text"], 300) for p in passages[:4]]
                ans      = await groq_answer(http, q, ctx_long)
            return ans

        answers = await asyncio.gather(*(handle(q) for q in questions))

    # 3. Persist log
    log = Data(
        user_query=json.dumps(questions),
        ai_response=json.dumps(answers),
        document_url=document.filename,
        timestamp=datetime.datetime.now(),
    )
    session.add(log)
    session.commit()

    return QueryResponse(answers=list(answers))
